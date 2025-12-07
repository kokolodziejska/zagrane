from docx import Document
from datetime import date
from io import BytesIO

def create_docx(data, comment, date):
    doc = Document()
    print(data)
    doc.add_paragraph(f'Proszę edytować do: {date}')
    table = doc.add_table(rows=1, cols=len(data['headers']))
    table.style = "Table Grid"

    # bold text
    hdr_cells = table.rows[0].cells
    for i, heading in enumerate(data['headers']):
        hdr_cells[i].text = heading
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Add remaining rows
    for row in data["rows"]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)

    doc.add_paragraph(comment)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


if __name__ == '__main__':
    data = [
        ["col1", 'col2', 'col3', 'col4', 'col5'],
        ['data1', 'data2', 'data3', 'data4', 'data5'],
        [1, 2, 3, 4, 5],
        ['odsao', ' ', 'dasdsadsa', '123', 12]
    ]

    comment = 'Proszę bardzo o szybkie wykonanie zadania!'
    today = date.today()
    bts = create_docx(data, comment, today)
    with open('test.docx', 'wb') as f:
        f.write(bts)
